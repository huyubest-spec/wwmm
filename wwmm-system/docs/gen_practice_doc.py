"""
项目实训文档生成器
基于项目实训文档模板生成《基于区块链的摄影作品投票存证系统》文档
"""
import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from copy import deepcopy

OUT = r'D:\GitHub\wwmm\wwmm-system\docs\项目实训文档-基于区块链的摄影作品投票存证系统的设计与实现.docx'
SS = r'D:\GitHub\wwmm\wwmm-system\screenshots'

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

# 默认样式
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_run_font(run, size=12, bold=False, color=None, font_name='宋体'):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_title(text, level=1, size=None, center=True):
    if size is None:
        size = {1: 18, 2: 16, 3: 14, 4: 12}.get(level, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, size=size, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_heading(text, level=1):
    sizes = {1: 18, 2: 16, 3: 14, 4: 13}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, size=sizes.get(level, 12), bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_para(text, size=12, bold=False, indent=True, line_space=1.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    pf = p.paragraph_format
    pf.line_spacing = line_space
    if indent:
        pf.first_line_indent = Pt(size * 2)
    pf.space_after = Pt(2)
    return p

def add_bullets(items, size=12):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(it)
        set_run_font(r, size=size)
        p.paragraph_format.line_spacing = 1.5

def add_table(headers, rows, widths=None, header_fill='D5E8F0'):
    n = len(headers)
    tbl = doc.add_table(rows=1, cols=n)
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=11, bold=True)
        # 设置表头底色
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), header_fill)
        tcPr.append(shd)
    # 设置列宽
    if widths:
        for i, w in enumerate(widths):
            for cell in tbl.columns[i].cells:
                cell.width = Cm(w)
    # 数据行
    for row in rows:
        cells = tbl.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(str(v))
            set_run_font(r, size=10.5)
    return tbl

def add_image(name, w=14, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    img_path = os.path.join(SS, name)
    if os.path.exists(img_path):
        r.add_picture(img_path, width=Cm(w))
    if caption:
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc = pc.add_run(caption)
        set_run_font(rc, size=10, bold=True)

def add_pagebreak():
    doc.add_page_break()

def add_code(text, lang=''):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Pt(0)

# ============== 封面 ==============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\n\n\n\n\n\n\n\n\n')
set_run_font(r, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('湖南科技职业学院')
set_run_font(r, size=20, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('毕业设计（项目实训）成果')
set_run_font(r, size=28, bold=True, color=(0x2c, 0x5f, 0xe0))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\n\n\n\n\n\n\n\n\n\n')
set_run_font(r, size=12)

# 信息表
add_table(
    ['选题名称', '基于区块链的摄影作品投票存证系统的设计与实现'],
    [
        ['选题类型', '方案设计类'],
        ['二级学院', '软件学院'],
        ['专业名称', '区块链技术应用'],
        ['班级名称', '区块链3232'],
        ['姓    名', '李文博'],
        ['学    号', '202103323201'],
        ['指导教师', '汪铭杰、罗斌'],
        ['完成时间', '2026年4月']
    ],
    widths=[3, 11]
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\n\n湖南科技职业学院教务处')
set_run_font(r, size=12, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('2024年11月')
set_run_font(r, size=12)

add_pagebreak()

# ============== 毕业设计真实性承诺及指导教师声明 ==============
add_title('毕业设计真实性承诺及指导教师声明', level=1)

add_heading('学生毕业设计真实性承诺', level=2)
add_para('本人郑重声明：所提交的毕业设计是本人在指导教师的指导下，独立进行研究工作所取得的成果，'
         '内容真实可靠，不存在抄袭、造假等学术不端行为。除设计方案中已经注明引用的内容外，'
         '本设计不含其他个人或集体已经发表或撰写过的研究成果。对本设计的研究做出重要贡献的个人和集体，'
         '均已在设计文档中明确注明。如被发现设计中存在抄袭、造假等学术不端行为，'
         '本人愿意承担相应的法律责任和一切后果。')

p = doc.add_paragraph()
r = p.add_run('\n\n学生（签名）：____________________')
set_run_font(r, size=12)

p = doc.add_paragraph()
r = p.add_run('2026年4月13日')
set_run_font(r, size=12)

add_heading('指导教师关于学生毕业设计真实性审核的声明', level=2)
add_para('本人郑重声明：已经对学生毕业设计所涉及的内容进行严格审核，确定其成果均由学生在本人指导下取得，'
         '对他人设计方案及成果的引用已经明确注明，不存在抄袭、造假等学术不端行为。')

p = doc.add_paragraph()
r = p.add_run('\n\n指导教师（签名）：____________________')
set_run_font(r, size=12)

p = doc.add_paragraph()
r = p.add_run('2026年4月15日')
set_run_font(r, size=12)

add_pagebreak()

# ============== 目录 ==============
add_title('目  录', level=1)
add_para('1  设计任务', size=12)
add_para('1.1 项目背景', size=12)
add_para('1.2 项目范围', size=12)
add_para('1.2.1 用户注册模块', size=12)
add_para('1.2.2 用户登录模块', size=12)
add_para('1.2.3 作品上传模块', size=12)
add_para('1.2.4 投票模块', size=12)
add_para('1.2.5 区块链存证模块', size=12)
add_para('1.3 项目风险分析', size=12)
add_para('1.4 任务分配', size=12)
add_para('1.5 项目所需资源', size=12)
add_para('2  设计思路与技术方案', size=12)
add_para('2.1 业务流程图', size=12)
add_para('2.2 系统用例图', size=12)
add_para('2.3 用例分析', size=12)
add_para('2.3.1 游客登录', size=12)
add_para('2.3.2 游客注册', size=12)
add_para('2.3.3 上传作品', size=12)
add_para('2.3.4 为作品投票', size=12)
add_para('2.3.5 校验作品存证', size=12)
add_para('2.3.15 管理员登录', size=12)
add_para('2.3.16 管理员审核', size=12)
add_para('2.3.17 区块链浏览器查询', size=12)
add_para('2.4 技术方案', size=12)
add_para('2.4.1 项目系统架构图', size=12)
add_para('2.4.2 数据库命名规则', size=12)
add_para('2.4.3 类命名规则', size=12)
add_para('2.4.4 方法、参数、成员变量、局部变量命名规则', size=12)
add_para('2.4.5 包名结构', size=12)
add_para('3  设计内容（过程）与说明', size=12)
add_para('3.1 区块链平台搭建与配置', size=12)
add_para('3.1.1 区块链底层平台结点拓扑设计', size=12)
add_para('3.1.2 中间件部署和平台配置', size=12)
add_para('3.1.3 区块链平台测试和验证', size=12)
add_para('3.2 智能合约设计与实现', size=12)
add_para('3.2.1 合约接口设计', size=12)
add_para('3.2.2 智能合约实现', size=12)
add_para('3.2.3 智能合约测试', size=12)
add_para('3.3 数据库设计', size=12)
add_para('3.3.1 数据库模型图', size=12)
add_para('3.3.2 数据表设计', size=12)
add_para('3.4 界面设计', size=12)
add_para('3.4.1 登录', size=12)
add_para('3.4.2 注册', size=12)
add_para('3.4.3 游客-作品列表', size=12)
add_para('3.4.4 游客-作品详情', size=12)
add_para('3.4.5 摄影师-我的作品', size=12)
add_para('3.4.6 摄影师-上传作品', size=12)
add_para('3.4.7 管理员-审核中心', size=12)
add_para('3.5 项目测试', size=12)
add_para('3.5.1 单元测试计划', size=12)
add_para('3.5.2 集成测试计划', size=12)
add_para('3.5.3 单元测试报告', size=12)
add_para('3.5.4 集成测试报告', size=12)
add_para('4  设计总结', size=12)
add_para('4.1 部署手册', size=12)
add_para('4.1.1 区块链平台搭建', size=12)
add_para('4.1.2 数据库初始化', size=12)
add_para('4.1.3 启动服务端', size=12)
add_para('4.1.4 启动客户端', size=12)
add_para('4.1.5 访问程序', size=12)
add_para('4.2 用户操作手册', size=12)
add_para('4.2.1 用户登录', size=12)
add_para('4.2.2 用户注册', size=12)
add_para('4.2.3 上传与投票功能', size=12)
add_para('4.3 总结感悟', size=12)
add_para('参考文献', size=12)

add_pagebreak()

# ============== 1 设计任务 ==============
add_title('1  设计任务', level=1)

add_heading('1.1 项目背景', level=2)
add_para('随着数字摄影的普及和互联网分享文化的兴起，摄影作品在网络空间被大量复制、传播、'
         '二次加工，作品归属与原创性证明变得日益困难。'
         '现有的内容平台大多使用中心化数据库来登记版权信息，'
         '数据库的"被篡改"、"被删除"、"被伪造"等风险依然存在，'
         '原作者难以低成本地证明"这张照片是我先拍的"。')

add_para('与此同时，公众对摄影作品的关注度持续上升，'
         '各类摄影赛事、摄影比赛、爱好者社区都需要一个公开可信的投票系统。'
         '传统的中心化投票系统存在刷票、暗箱操作、数据可被运维方任意修改等问题，'
         '投票结果难以让所有参与者信服。')

add_para('从技术层面来看，区块链技术具有去中心化、公开透明、不可篡改、可追溯等核心特性，'
         '可天然解决上述两个痛点：')

add_bullets([
    '将摄影作品的 SHA-256 哈希连同作者、拍摄时间、描述等元数据写入区块链，'
    '上链后任何人无法篡改，相当于给作品颁发了一张"数字身份证"。',
    '将每一次投票作为交易写入区块链，所有投票记录公开可查，'
    '任何人都可以独立验证票数，投票结果不可被运维方暗箱操作。',
    '区块链浏览器让普通用户也能低门槛地查询链上数据，提升了系统的公信力。'
])

add_para('因此，本毕业设计《基于区块链的摄影作品投票存证系统》旨在构建一个'
         '集作品存证、人气投票、链上验证于一体的去中心化平台，'
         '让摄影作品从上传那一刻起便拥有不可伪造的存在证明，'
         '让投票从那一刻起便无可争议。')

add_heading('1.2 项目范围', level=2)
add_para('根据用户需求，本项目主要分为 7 个主要模块：'
         '用户注册模块、用户登录模块、作品上传模块、'
         '管理员审核模块、投票模块、区块链存证模块、'
         '区块链浏览器模块。')

add_heading('1.2.1 用户注册模块', level=3)
add_para('本平台的注册模块分为两大角色：普通用户、摄影师，不同角色注册流程不同。')

add_para('（1）普通用户注册')
add_para('普通用户注册时需要填写账号、密码、真实姓名、性别、手机号、邮箱等。'
         '注册成功后信息将会添加到数据库中，并同时写入区块链，'
         '用户获得投票权限。')

add_para('（2）摄影师注册')
add_para('摄影师注册时需要填写账号、密码、真实姓名、手机号、邮箱、个人简介等。'
         '注册成功后信息将会添加到数据库中，'
         '并同时写入区块链，摄影师获得作品上传权限。')

add_heading('1.2.2 用户登录模块', level=3)
add_para('已注册用户通过账号 + 密码登录，登录成功后服务端返回会话 Token，'
         '前端将 Token 写入 localStorage，'
         '后续请求通过 Token 鉴权。')

add_heading('1.2.3 作品上传模块', level=3)
add_para('摄影师可上传 JPG/PNG 格式的图片，'
         '填写标题、描述、分类、拍摄地点、拍摄时间、器材信息等元数据。'
         '上传时服务端自动计算图片的 SHA-256 哈希，'
         '哈希与元数据一同作为"作品存证"交易写入区块链。')

add_heading('1.2.4 投票模块', level=3)
add_para('注册用户对已审核通过的作品进行投票，每位用户对同一作品仅可投 1 票。'
         '投票记录作为"投票存证"交易写入区块链。'
         '作品详情页实时显示链上累计票数。')

add_heading('1.2.5 区块链存证模块', level=3)
add_para('采用自研 PoW 区块链引擎，将所有作品存证、投票记录打包进区块，'
         '每个区块包含：')
add_bullets([
    '区块高度（Index）',
    '前一区块哈希（PrevHash）',
    'Merkle 根（MerkleRoot）',
    '时间戳（Timestamp）',
    '工作量证明随机数（Nonce）',
    '挖矿难度（Difficulty）',
    '本区块哈希（Hash）',
    '交易数（TxCount）',
    '打包者（Miner）'
])

add_heading('1.3 项目风险分析', level=2)
add_para('在本项目中，可能出现需求不准、技术路线变更等项目风险，具体如下表。')
add_table(
    ['序号', '风险', '概率（%）', '严重性（%）', '措施'],
    [
        ['1', '页面展示出错', '70', '60', '多浏览器调试，编写响应式 CSS'],
        ['2', '服务器启动错误', '60', '70', '分步启动，先数据库后 Go 服务'],
        ['3', '写功能报错', '80', '70', '仔细写代码，编译通过再联调'],
        ['4', '代码混乱', '50', '100', '完成一步，整理一步，按 MVC 分层'],
        ['5', '功能需求错误', '20', '50', '多次分析需求，分模块迭代'],
        ['6', '网站抗压能力差', '20', '100', '可使用负载均衡、缓存技术'],
        ['7', '区块链哈希冲突', '10', '80', '采用 SHA-256 + 唯一索引防重']
    ]
)

add_heading('1.4 任务分配', level=2)
add_para('本项目由个人独立完成，所有任务都是本人分配以及完成。')
add_table(
    ['序号', '任务', '负责人'],
    [
        ['1', '需求分析与文档撰写', '李文博'],
        ['2', '区块链平台搭建（自研引擎）', '李文博'],
        ['3', '智能合约（Solidity）设计与实现', '李文博'],
        ['4', '数据库设计与初始化', '李文博'],
        ['5', '后端服务接口实现（Go + Gin）', '李文博'],
        ['6', '前端页面设计与实现（Vue 2）', '李文博'],
        ['7', '项目整合与联调', '李文博'],
        ['8', '测试', '李文博'],
        ['9', 'PPT 与报告撰写', '李文博']
    ]
)

add_heading('1.5 项目所需资源', level=2)
add_para('本项目主要所用到的工具及版本如下表所示。')
add_table(
    ['序号', '资源', '作用', '占用时间', '可用状态', '获得途径'],
    [
        ['1', 'Visual Studio Code 1.87.1', '编写前后端代码', '贯穿整个项目', '可用', '网上下载'],
        ['2', 'GoLand 2022.2.4 / Go 1.21.5', '编写后端代码', '贯穿整个项目', '可用', '网上下载'],
        ['3', 'Node.js v24.13.1 / Vite 4.5', '前端构建工具', '贯穿整个项目', '可用', '网上下载'],
        ['4', 'MySQL 8.0.45', '业务数据存储', '贯穿整个项目', '可用', '本机已装'],
        ['5', 'WPS Office', '编写文档', '文档撰写阶段', '可用', '网上下载'],
        ['6', 'Navicat 12.0.18', '数据库管理', '调试阶段', '可用', '网上下载'],
        ['7', 'Solidity 0.4.25', '智能合约开发', '合约设计阶段', '可用', 'npm 安装'],
        ['8', 'Chrome 浏览器', '页面调试与截图', '测试阶段', '可用', '本机已装']
    ]
)

add_pagebreak()

# ============== 2 设计思路与技术方案 ==============
add_title('2  设计思路与技术方案', level=1)
add_para('经过分析与调研，本项目确定了 7 个主要功能模块：'
         '用户注册模块、用户登录模块、作品上传模块、管理员审核模块、'
         '投票模块、区块链存证模块、区块链浏览器模块。'
         '技术栈采用主流框架完成：前端 Vue 2.0 + Element UI，'
         '后端 Go 1.21 + Gin，数据库 MySQL 8.0，'
         '区块链平台自研（PoW + Merkle 树 + 链式结构），'
         '智能合约采用 Solidity 0.4.25。')

add_heading('2.1 业务流程图', level=2)
add_para('（1）摄影师上传作品并触发存证上链')
add_para('当摄影师在前端填写作品信息并提交图片后，'
         '后端首先将图片保存到本地 uploads 目录，'
         '同时计算 SHA-256 哈希。'
         '然后在数据库中创建 photo 记录（status=0 待审核），'
         '并在 tx 表中插入一笔 tx_type=1 的"作品存证"交易，'
         '调用 PackPendingTxs() 立即打包为一个新区块写入数据库，'
         '同步更新 photo.is_on_chain=1 与 chain_tx_hash。')

add_para('（2）管理员审核作品')
add_para('管理员登录后访问 /#/admin，'
         '前端调用 /api/photo/pending 获取待审核列表。'
         '管理员点击"通过"或"拒绝"并填写审核意见，'
         '前端调用 /api/photo/:id/audit。'
         '后端更新 photo.status，并将审核记录写入 photo_audit_log。')

add_para('（3）用户投票并触发投票上链')
add_para('已登录用户在作品详情页点击"为它投一票"，'
         '前端调用 /api/photo/:id/vote。'
         '后端校验（作品已审核、非本人、未投过票），'
         '在 vote 表中插入投票记录，'
         '在 tx 表中插入 tx_type=2 的"投票存证"交易，'
         '调用 PackPendingTxs() 立即打包上链。'
         '前端跳转至区块浏览器即可查看该笔交易。')

add_heading('2.2 系统用例图', level=2)
add_para('经分析，本项目涉及三大类角色：游客 / 普通用户、摄影师、管理员。'
         '其中普通用户可注册、登录、浏览作品、为作品投票、'
         '查询作品存证信息、查看区块链浏览器；'
         '摄影师在普通用户的基础上可上传作品、查看我的作品；'
         '管理员可审核待上架作品、查看全网链上数据、'
         '管理用户。')

add_heading('2.3 用例分析', level=3)
add_para('从角色划分来看，普通用户涉及的用例共计 6 个、摄影师涉及的用例共计 4 个、'
         '管理员涉及的用例共计 3 个。下面就核心用例的基本路径、'
         '扩展路径和异常事件流，以及前置条件与后置条件设计如下：')

add_heading('2.3.1 游客登录', level=3)
add_table(
    ['内容', '说明'],
    [
        ['用例编号', '001'],
        ['用例名称', '登录'],
        ['用例说明', '用户输入账号和密码，进入平台，进行相应的权限操作。'],
        ['参与者', '游客 / 已注册用户'],
        ['前置条件', '该账户已存在'],
        ['后置条件', '登录成功后，跳转到首页'],
        ['基本路径', '（1）进入平台点击登录按钮（2）输入账号和密码（3）点击登录（4）系统校验提交的信息是否正确（5）登录成功'],
        ['扩展路径', '无'],
        ['异常事件流', '账号或密码不正确，提示用户名或密码错误'],
        ['补充说明', '没有进行登录验证不能进行投票、上传等操作']
    ]
)

add_heading('2.3.2 游客注册', level=3)
add_table(
    ['内容', '说明'],
    [
        ['用例编号', '002'],
        ['用例名称', '用户注册'],
        ['用例说明', '用户输入账号和密码以及个人信息，注册平台用户'],
        ['参与者', '游客'],
        ['前置条件', '该账号未被注册'],
        ['后置条件', '注册成功后跳转到登录界面'],
        ['基本路径', '（1）进入平台点击注册按钮（2）选择注册为普通用户或摄影师（3）输入账号和密码以及个人信息（4）点击注册（5）系统校验账号是否被注册（6）注册成功'],
        ['扩展路径', '（1）点击注册按钮（2）仅输入必填项（3）点击注册'],
        ['异常事件流', '账号已被注册，提示请使用其他账号'],
        ['补充说明', '无']
    ]
)

add_heading('2.3.3 上传作品', level=3)
add_table(
    ['内容', '说明'],
    [
        ['用例编号', '003'],
        ['用例名称', '上传作品'],
        ['用例说明', '摄影师上传摄影作品并触发区块链存证'],
        ['参与者', '摄影师'],
        ['前置条件', '摄影师已登录'],
        ['后置条件', '跳转到"我的作品"页面'],
        ['基本路径', '（1）进入首页（2）点击上传作品（3）选择图片并填写元数据（4）点击提交（5）后端计算 SHA-256 并写入数据库（6）后端生成存证交易并立即打包上链（7）系统提示上传成功'],
        ['扩展路径', '无'],
        ['异常事件流', '图片格式不支持 / 大小超过 10MB'],
        ['补充说明', '重复上传同一图片会被拒绝（哈希唯一索引）']
    ]
)

add_heading('2.3.4 为作品投票', level=3)
add_table(
    ['内容', '说明'],
    [
        ['用例编号', '004'],
        ['用例名称', '为作品投票'],
        ['用例说明', '注册用户对已审核通过的作品投票'],
        ['参与者', '注册用户'],
        ['前置条件', '用户已登录，作品状态为"已通过"'],
        ['后置条件', '投票记录写入区块链，作品得票数 +1'],
        ['基本路径', '（1）进入作品详情页（2）点击"为它投一票"（3）后端校验（4）写入 vote 表（5）生成投票存证交易（6）打包上链（7）前端显示交易哈希'],
        ['扩展路径', '无'],
        ['异常事件流', '已投过票 / 作品未审核 / 给自己投票'],
        ['补充说明', '每用户每作品仅 1 票']
    ]
)

add_heading('2.3.5 校验作品存证', level=3)
add_table(
    ['内容', '说明'],
    [
        ['用例编号', '005'],
        ['用例名称', '校验作品存证'],
        ['用例说明', '输入 SHA-256 哈希查询其是否已上链'],
        ['参与者', '任意访问者'],
        ['前置条件', '用户知道某图片的 SHA-256 哈希'],
        ['后置条件', '展示该哈希在链上的存证信息'],
        ['基本路径', '（1）进入哈希校验页（2）输入 64 字符哈希（3）点击校验（4）系统查询 photo 表（5）返回作品及链交易信息'],
        ['扩展路径', '无'],
        ['异常事件流', '哈希格式不正确 / 该哈希在链上未找到'],
        ['补充说明', '用于第三方独立验证作品归属']
    ]
)

add_heading('2.3.15 管理员登录', level=3)
add_table(
    ['内容', '说明'],
    [
        ['用例编号', '015'],
        ['用例名称', '管理员登录'],
        ['用例说明', '管理员输入账号和密码进入后台'],
        ['参与者', '管理员'],
        ['前置条件', '账号已存在且 role=2'],
        ['后置条件', '跳转到审核中心'],
        ['基本路径', '（1）进入平台点击登录（2）输入管理员账号密码（3）点击登录（4）系统校验角色'],
        ['扩展路径', '无'],
        ['异常事件流', '账号或密码错误'],
        ['补充说明', '普通用户无权限访问审核中心']
    ]
)

add_heading('2.3.16 管理员审核', level=3)
add_table(
    ['内容', '说明'],
    [
        ['用例编号', '016'],
        ['用例名称', '管理员审核'],
        ['用例说明', '审核摄影师提交的作品'],
        ['参与者', '管理员'],
        ['前置条件', '已有待审核作品'],
        ['后置条件', '作品状态变为已通过或已拒绝'],
        ['基本路径', '（1）进入审核中心（2）查看待审核列表（3）点击"通过"或"拒绝"并填写意见（4）提交（5）系统更新状态并写入审计日志'],
        ['扩展路径', '无'],
        ['异常事件流', '无'],
        ['补充说明', '已通过作品允许投票，已拒绝作品隐藏']
    ]
)

add_heading('2.3.17 区块链浏览器查询', level=3)
add_table(
    ['内容', '说明'],
    [
        ['用例编号', '017'],
        ['用例名称', '区块链浏览器查询'],
        ['用例说明', '查看链上所有区块、交易、状态'],
        ['参与者', '任意访问者'],
        ['前置条件', '无'],
        ['后置条件', '展示链数据'],
        ['基本路径', '（1）进入区块链浏览器（2）查看最新区块（3）点击区块查看详情（4）点击交易查看 payload（5）跳转到哈希校验页'],
        ['扩展路径', '无'],
        ['异常事件流', '无'],
        ['补充说明', '无']
    ]
)

add_heading('2.4 技术方案', level=2)
add_para('本"基于区块链的摄影作品投票存证系统"项目采用主流框架技术完成。'
         '前端使用的框架主要有 Vue 2.7 + Element UI 2.15 + Axios，'
         '后端使用技术主要有 Go 1.21 + Gin 1.9 + database/sql，'
         '数据库采用 MySQL 8.0.45 系统，'
         '区块链平台自研实现（SHA-256 + PoW + Merkle 树 + 链式结构），'
         '智能合约采用 Solidity 0.4.25 编写。'
         '开发过程中，按 MVC 分层组织代码，'
         '数据库与包命名遵守统一规范。')

add_heading('2.4.1 项目系统架构图', level=3)
add_para('整个系统采用经典的前后端分离架构：')
add_bullets([
    '【前端层】Vue 2 SPA + Vue Router + Element UI + Axios',
    '【API 网关层】Gin HTTP Server（监听 :8080），统一处理 CORS、鉴权、路由',
    '【业务层】Service 层（user / photo / vote / chain），封装业务规则',
    '【数据访问层】DAO 层（user / photo / vote / chain），通过 database/sql 操作 MySQL',
    '【区块链层】自研 blockchain/ 包（merkle.go / block.go），提供 PoW 挖矿、链打包、链查询',
    '【存储层】MySQL 8.0（业务数据 + 区块数据 + 交易数据）',
    '【文件存储】本地 uploads 目录（图片二进制）'
])

add_heading('2.4.2 数据库命名规则', level=3)
add_para('表达是与否概念的字段，必须使用 is_xxx 的方式命名，数据类型是 unsigned tinyint'
         '（1 表示是，0 表示否）。说明：任何字段如果为非负数，必须是 unsigned。')
add_para('正例：表达逻辑删除的字段名 is_deleted，1 表示删除，0 表示未删除。')
add_para('表名、字段名必须使用小写字母或数字，禁止出现数字开头，'
         '禁止两个下划线中间只出现数字。')
add_para('正例：aliyun_admin，rdc_config，level3_name')
add_para('反例：AliyunAdmin，rdcConfig，level_3_name')
add_para('表名不使用复数名词。表名应该仅仅表示表里面的实体内容，'
         '不应该表示实体数量，对应于 DO 类名也是单数形式，符合表达习惯。')

add_heading('2.4.3 类命名规则', level=3)
add_para('类名使用 UpperCamelCase 风格，必须遵从驼峰形式，'
         '但以下情形例外：DO / BO / DTO / VO / AO')
add_para('正例：MarcoPolo / UserDO / XmlService / TcpUdpDeal / PhotoEvidence')
add_para('反例：macroPolo / UserDo / XMLService / TCPUDPDeal')

add_heading('2.4.4 方法、参数、成员变量、局部变量命名规则', level=3)
add_para('方法名、参数名、成员变量、局部变量都统一使用 lowerCamelCase 风格，'
         '必须遵从驼峰形式。')
add_para('正例：localValue / getHttpMessage() / inputUserId / packBlock() / castVote()')

add_heading('2.4.5 包名结构', level=3)
add_bullets([
    'Controller/* —— HTTP 控制器层',
    'Service/*    —— 业务逻辑层',
    'Dao/*        —— 与数据库交互的 SQL',
    'Model/*      —— 与数据库字段对应的实体对象',
    'Blockchain/* —— 自研区块链引擎（merkle.go / block.go）',
    'Config/*     —— 配置文件',
    'Utils/*      —— 工具函数（crypto / auth / response / db）',
    'main.go      —— 项目入口函数',
    'seed.go      —— 种子用户初始化',
    'contracts/   —— Solidity 智能合约设计稿'
])

add_pagebreak()

# ============== 3 设计内容（过程）与说明 ==============
add_title('3  设计内容（过程）与说明', level=1)
add_para('《基于区块链的摄影作品投票存证系统》设计完成后，'
         '交付内容主要有：数据库设计、类设计与各业务时序图设计，'
         '以及界面设计。本系统涉及的功能点共计 22 个，'
         '数据库表共计 8 张，主要业务类共计 6 个，各类型界面共计 11 个。')

add_heading('3.1 区块链平台搭建与配置', level=2)

add_heading('3.1.1 区块链底层平台结点拓扑设计', level=3)
add_para('考虑到本地 Windows 环境部署 FISCO BCOS 4 节点联盟链集群复杂度较高，'
         '本项目使用 Go 语言在 backend/blockchain/ 目录下完整复现了一个轻量级 PoW 区块链。'
         '该引擎实现了与 FISCO BCOS 一致的链式结构、Merkle 根、交易验证等核心特性，'
         '同时具备轻量、可移植、易于理解的优势。')

add_para('结点拓扑设计：')
add_table(
    ['序号', '类型', '网络地址', '描述'],
    [
        ['1', 'Host', '127.0.0.1', 'CPU: i5-1135G7 Memory: 16G 本机'],
        ['2', 'Node', 'wwmm-node-01', 'Go 自研 PoW 节点'],
        ['3', 'Database', 'MySQL 8.0', '存储区块、交易、状态'],
        ['4', 'API', ':8080', 'RESTful API 服务'],
        ['5', 'Frontend', 'Vue 2 SPA', '区块链浏览器 UI']
    ]
)

add_heading('3.1.2 中间件部署和平台配置', level=3)
add_para('启动后端 Go 服务：')
add_code('cd backend/\ngo build -o wwmm-server.exe .\n./wwmm-server.exe\n# 输出:\n# [OK] 数据库连接成功\n# [OK] WWMM 后端启动于 :8080')

add_para('启动前端 Vite 服务：')
add_code('cd frontend/\nnpm install\nnpm run dev\n# 输出:\n# VITE v4.5.14  ready in 586 ms\n# ➜  Local:   http://localhost:5173/')

add_heading('3.1.3 区块链平台测试和验证', level=3)
add_para('启动后通过以下方式验证区块链引擎工作：')
add_bullets([
    '访问 http://localhost:8080/api/chain/state 查看链状态',
    '访问 http://localhost:5173/#/chain 进入区块链浏览器',
    '上传一张图片后查看 /api/chain/txs 是否新增"作品存证"交易',
    '投票后查看 /api/chain/txs 是否新增"投票存证"交易',
    '通过 /api/chain/block/:index 查看区块详情，验证哈希链、Merkle 根、Nonce 正确'
])

add_heading('3.2 智能合约设计与实现', level=2)
add_para('本项目智能合约作为"设计稿"保存于 contracts/ 目录，'
         '展示了如果部署到 FISCO BCOS 联盟链时合约应该如何实现。'
         '本项目在 Go 端以等价逻辑实现了相同的状态机与事件，'
         '保证业务行为一致。')

add_heading('3.2.1 合约接口设计', level=3)

add_para('a. Roles 角色库')
add_para('提供 RBAC 角色管理能力，被其他合约复用。关键方法：')
add_bullets([
    'Roles.add(Role, address) — 授予角色',
    'Roles.remove(Role, address) — 撤销角色',
    'Roles.has(Role, address) — 查询是否拥有该角色'
])

add_para('b. Admin 合约')
add_para('提供平台管理员角色管理。关键方法：')
add_bullets([
    'addAdmin(address) — 管理员可新增其他管理员',
    'renounceAdmin() — 管理员可放弃自己权限',
    'isAdmin(address) — 查询是否是管理员'
])

add_para('c. Photographer 合约（继承自 Admin）')
add_para('提供摄影师注册与管理。关键方法：')
add_bullets([
    'register(string realName, string phone) — 注册成为摄影师',
    'removePhotographer(address) — 管理员可注销摄影师',
    'getInfo(address) — 查询摄影师信息',
    '_incrementPhotoCount(address) — 上传作品后自增计数'
])

add_para('d. PhotoEvidence 合约（继承自 Photographer）')
add_para('提供作品存证 + 投票核心逻辑。关键方法：')
add_bullets([
    'submitPhoto(title, imageHash, description, category, shootLocation) — 提交作品存证',
    'auditPhoto(photoId, approve, comment) — 管理员审核',
    'voteFor(photoId) — 为作品投票',
    'getPhoto(photoId) — 查询作品详情',
    'getPhotoByHash(imageHash) — 通过图片哈希反查',
    'getRanking(topN) — 获取排行榜'
])

add_heading('3.2.2 智能合约实现', level=3)
add_para('合约核心代码片段（完整代码见 contracts/PhotoEvidence.sol）：')
add_code('''// SPDX-License-Identifier: MIT
pragma solidity ^0.4.25;
import "./Photographer.sol";

contract PhotoEvidence is Photographer {

    enum PhotoStatus { Pending, Approved, Rejected }

    struct Photo {
        uint256       photoId;
        address       photographer;
        string        imageHash;        // 64 字符 SHA-256 十六进制
        string        title;
        string        description;
        PhotoStatus   status;
        uint256       voteCount;
    }

    struct Vote { address voter; uint256 photoId; uint256 voteTime; bool valid; }

    uint256 private _photoSeq;
    mapping(uint256 => Photo) private _photos;
    mapping(uint256 => Vote[]) private _votes;
    mapping(address => uint256[]) private _userVotes;
    mapping(bytes32 => uint256) private _hashIndex;

    event PhotoSubmitted(uint256 indexed photoId, address indexed photographer, string imageHash);
    event PhotoAudited(uint256 indexed photoId, PhotoStatus status);
    event Voted(uint256 indexed photoId, address indexed voter, uint256 voteCount);

    function submitPhoto(string title, string imageHash, string description,
                        string category, string shootLocation)
        public onlyPhotographer returns (uint256)
    {
        bytes32 h = keccak256(abi.encodePacked(imageHash));
        require(_hashIndex[h] == 0, "photo already submitted");
        require(bytes(title).length > 0, "title required");
        require(bytes(imageHash).length == 64, "image hash must be 64 hex chars");

        _photoSeq += 1;
        uint256 pid = _photoSeq;
        _photos[pid] = Photo({
            photoId: pid, photographer: msg.sender, title: title,
            imageHash: imageHash, description: description,
            status: PhotoStatus.Pending, voteCount: 0
        });
        _hashIndex[h] = pid;
        _incrementPhotoCount(msg.sender);
        emit PhotoSubmitted(pid, msg.sender, imageHash);
        return pid;
    }

    function voteFor(uint256 photoId) public photoExists(photoId) {
        Photo storage p = _photos[photoId];
        require(p.status == PhotoStatus.Approved, "not approved");
        require(p.photographer != msg.sender, "cannot vote for yourself");
        require(!_hasVoted(msg.sender, photoId), "already voted");

        _votes[photoId].push(Vote(msg.sender, photoId, now, true));
        _userVotes[msg.sender].push(photoId);
        p.voteCount += 1;
        emit Voted(photoId, msg.sender, p.voteCount);
    }
}''')

add_heading('3.2.3 智能合约测试', level=3)
add_para('由于本地 Windows 环境未部署 FISCO BCOS 节点，'
         '合约的等价测试通过 Go 端的 Service 层与 DAO 层完成。'
         '本项目的所有合约接口都有对应 Go 实现，行为与合约完全一致：')

add_table(
    ['合约方法', 'Go 等价实现', '测试结果'],
    [
        ['submitPhoto', 'service.SubmitPhoto()', 'PASS'],
        ['auditPhoto', 'service.AuditPhoto()', 'PASS'],
        ['voteFor', 'service.CastVote()', 'PASS'],
        ['getPhoto', 'service.GetPhotoDetail()', 'PASS'],
        ['getPhotoByHash', 'dao.GetPhotoByHash()', 'PASS'],
        ['getRanking', 'Ranking.vue 排序逻辑', 'PASS']
    ]
)

add_heading('3.3 数据库设计', level=2)
add_para('数据库设计是系统设计的基础。本项目数据库 wwmm_db，'
         '一共包含 8 张数据表，分别为用户表、摄影作品表、投票记录表、'
         '区块表、交易表、审核日志表、链状态表、会话表。'
         '表与表之间的关系以及各表具体设计如下。')

add_heading('3.3.1 数据库模型图', level=3)
add_para('数据库主要表关系：')
add_bullets([
    '【user 1 — N photo】一位摄影师可上传多张作品',
    '【photo 1 — N vote】一张作品可被多用户投票',
    '【photo 1 — N photo_audit_log】一张作品可有多次审核记录',
    '【block 1 — N tx】一个区块包含多笔交易',
    '【user 1 — N session】一个用户可有多端登录',
    '【chain_state 1】固定 1 条记录，记录最新区块'
])

add_heading('3.3.2 数据表设计', level=3)

add_para('（1）用户表 user')
add_table(
    ['字段名', '字段说明', '类型', '长度', '允许空', '备注'],
    [
        ['user_id', '用户ID', 'INT', '11', 'NO', '主键 自增'],
        ['username', '登录账号', 'VARCHAR', '64', 'NO', '唯一'],
        ['password_hash', '密码哈希', 'VARCHAR', '255', 'NO', 'SHA-256(salt:password)'],
        ['salt', '盐值', 'VARCHAR', '32', 'NO', '随机 16 字节'],
        ['phone', '手机号', 'VARCHAR', '20', 'YES', ''],
        ['email', '邮箱', 'VARCHAR', '128', 'YES', ''],
        ['real_name', '真实姓名', 'VARCHAR', '64', 'YES', ''],
        ['sex', '性别', 'TINYINT', '', 'YES', '0-未知 1-男 2-女'],
        ['avatar', '头像', 'VARCHAR', '255', 'YES', ''],
        ['bio', '简介', 'VARCHAR', '255', 'YES', ''],
        ['role', '角色', 'TINYINT', '', 'NO', '0-普通 1-摄影师 2-管理员'],
        ['status', '状态', 'TINYINT', '', 'NO', '0-禁用 1-启用'],
        ['created_at', '创建时间', 'DATETIME', '', 'NO', 'CURRENT_TIMESTAMP'],
        ['updated_at', '更新时间', 'DATETIME', '', 'NO', 'ON UPDATE']
    ]
)

add_para('（2）摄影作品表 photo')
add_table(
    ['字段名', '字段说明', '类型', '长度', '允许空', '备注'],
    [
        ['photo_id', '作品ID', 'INT', '11', 'NO', '主键 自增'],
        ['title', '标题', 'VARCHAR', '128', 'NO', ''],
        ['description', '描述', 'TEXT', '', 'YES', ''],
        ['image_url', '图片URL', 'VARCHAR', '255', 'NO', '相对路径'],
        ['image_hash', '图片SHA-256', 'VARCHAR', '64', 'NO', '唯一 上链存证'],
        ['file_size', '文件大小', 'INT', '', 'NO', '字节'],
        ['photographer_id', '摄影师ID', 'INT', '11', 'NO', '外键→user.user_id'],
        ['category', '分类', 'VARCHAR', '32', 'YES', ''],
        ['shoot_location', '拍摄地点', 'VARCHAR', '128', 'YES', ''],
        ['shoot_time', '拍摄时间', 'DATE', '', 'YES', ''],
        ['camera_info', '器材', 'VARCHAR', '128', 'YES', ''],
        ['status', '状态', 'TINYINT', '', 'NO', '0-待审核 1-通过 2-拒绝'],
        ['audit_comment', '审核意见', 'VARCHAR', '255', 'YES', ''],
        ['vote_count', '得票数', 'INT', '', 'NO', '默认 0'],
        ['view_count', '浏览数', 'INT', '', 'NO', '默认 0'],
        ['is_on_chain', '是否上链', 'TINYINT', '', 'NO', '0-否 1-是'],
        ['chain_tx_hash', '上链交易哈希', 'VARCHAR', '64', 'YES', ''],
        ['created_at', '创建时间', 'DATETIME', '', 'NO', ''],
        ['updated_at', '更新时间', 'DATETIME', '', 'NO', '']
    ]
)

add_para('（3）投票记录表 vote')
add_table(
    ['字段名', '字段说明', '类型', '长度', '允许空', '备注'],
    [
        ['vote_id', '投票ID', 'INT', '11', 'NO', '主键 自增'],
        ['user_id', '用户ID', 'INT', '11', 'NO', '外键→user'],
        ['photo_id', '作品ID', 'INT', '11', 'NO', '外键→photo'],
        ['tx_hash', '投票交易哈希', 'VARCHAR', '64', 'YES', '唯一'],
        ['created_at', '投票时间', 'DATETIME', '', 'NO', 'CURRENT_TIMESTAMP'],
        ['UNIQUE KEY', 'uk_user_photo', '', '', '', 'user_id+photo_id 唯一']
    ]
)

add_para('（4）区块表 block')
add_table(
    ['字段名', '字段说明', '类型', '长度', '允许空', '备注'],
    [
        ['block_id', '区块ID', 'INT', '11', 'NO', '主键 自增'],
        ['index_num', '区块高度', 'INT', '11', 'NO', '从 1 开始'],
        ['prev_hash', '前一区块哈希', 'VARCHAR', '64', 'NO', ''],
        ['merkle_root', 'Merkle根', 'VARCHAR', '64', 'NO', ''],
        ['timestamp', '时间戳', 'BIGINT', '', 'NO', '秒'],
        ['nonce', '工作量证明随机数', 'BIGINT', '', 'NO', ''],
        ['difficulty', '挖矿难度', 'INT', '', 'NO', '默认 4'],
        ['hash', '本区块哈希', 'VARCHAR', '64', 'NO', '唯一'],
        ['tx_count', '交易数', 'INT', '', 'NO', '默认 0'],
        ['miner', '打包者', 'VARCHAR', '128', 'YES', ''],
        ['created_at', '入库时间', 'DATETIME', '', 'NO', '']
    ]
)

add_para('（5）交易表 tx')
add_table(
    ['字段名', '字段说明', '类型', '长度', '允许空', '备注'],
    [
        ['tx_id', '交易ID', 'INT', '11', 'NO', '主键 自增'],
        ['tx_hash', '交易哈希', 'VARCHAR', '64', 'NO', '唯一'],
        ['block_id', '所在区块ID', 'INT', '11', 'YES', '未打包前为 NULL'],
        ['tx_type', '交易类型', 'TINYINT', '', 'NO', '1-作品存证 2-投票'],
        ['sender', '发送方', 'VARCHAR', '128', 'NO', '用户账号'],
        ['payload', '载荷 JSON', 'TEXT', '', 'NO', ''],
        ['status', '状态', 'TINYINT', '', 'NO', '0-待打包 1-已打包 2-失败'],
        ['created_at', '创建时间', 'DATETIME', '', 'NO', '']
    ]
)

add_heading('3.4 界面设计', level=2)
add_para('本平台页面共计 11 张，页面设计如下。')

add_heading('3.4.1 登录', level=3)
add_para('用户在浏览器输入网址后，系统将跳转到登录页，要求用户输入有效登录信息。')
add_image('02_login.png', w=12, caption='图 3.1  登录页面')

add_heading('3.4.2 注册', level=3)
add_para('用户点击右上角"免费注册"可跳转到注册界面。')
add_image('03_register.png', w=12, caption='图 3.2  注册界面')

add_heading('3.4.3 游客-作品列表', level=3)
add_para('用户可访问首页浏览所有已审核通过的作品。')
add_para('首页展示链状态卡片（最新区块高度、累计交易、作品存证数、投票记录数）、'
         '热门作品瀑布流。')
add_image('01_home.png', w=15, caption='图 3.3  首页-作品广场')

add_heading('3.4.4 游客-作品详情', level=3)
add_para('用户点击作品可进入详情页，查看作品图片、标题、作者、'
         '拍摄信息、链上存证信息。已登录用户可投票。')
add_image('12_photo_detail_voter.png', w=15, caption='图 3.4  作品详情页')

add_heading('3.4.5 摄影师-我的作品', level=3)
add_para('摄影师登录后可查看自己上传的所有作品，'
         '以及每张作品的状态（待审核/已通过/已拒绝）、是否已上链、得票数。')
add_image('10_my_photos.png', w=15, caption='图 3.5  摄影师-我的作品')

add_heading('3.4.6 摄影师-上传作品', level=3)
add_para('摄影师在上传页填写标题、描述、分类等元数据，'
         '拖拽或选择图片后点击"提交存证"，'
         '后端立即计算 SHA-256 并打包上链。')
add_image('11_upload.png', w=15, caption='图 3.6  上传作品页')

add_heading('3.4.7 管理员-审核中心', level=3)
add_para('管理员可查看所有待审核作品，对作品点击"通过"或"拒绝"，'
         '并填写审核意见。审核通过后作品才能被投票。')
add_image('09_admin_audit.png', w=15, caption='图 3.7  审核中心')

add_para('其他关键页面：')
add_image('04_ranking.png', w=15, caption='图 3.8  人气排行榜')
add_image('05_chain_explorer.png', w=15, caption='图 3.9  区块链浏览器')
add_image('06_chain_block_detail.png', w=15, caption='图 3.10  区块详情')
add_image('07_chain_tx_detail.png', w=15, caption='图 3.11  交易详情')
add_image('08_chain_verify.png', w=15, caption='图 3.12  哈希存证校验')

add_heading('3.5 项目测试', level=2)
add_para('为了保证系统能够顺利上线运行，'
         '实现过程中，就系统已实现的功能是否达到任务目标、'
         '业务流程是否正确、系统运行是否稳定、'
         '缺陷率是否在可接受范围进行了测试。')

add_heading('3.5.1 单元测试计划', level=3)
add_table(
    ['项目', '内容'],
    [
        ['单元测试范围', '各功能模块'],
        ['单元测试方法', '人工输入值 + Go 单元函数 + curl 接口验证'],
        ['单元测试环境', 'Go 1.21.5, Node.js v24, MySQL 8.0.45, Windows 11'],
        ['将产生的文档', '单元测试用例，测试报告'],
        ['单元测试任务/优先级', '见下表']
    ]
)

add_table(
    ['要测试的功能', '预期结果', '进度日期', '实际日期', '负责人'],
    [
        ['用户注册', '能否成功注册账号', '2026-03-25', '2026-03-25', '李文博'],
        ['用户登录', '能否正常登录', '2026-03-25', '2026-03-25', '李文博'],
        ['退出登录', '能否成功退出', '2026-03-26', '2026-03-26', '李文博'],
        ['图片哈希计算', 'SHA-256 是否正确', '2026-03-26', '2026-03-26', '李文博'],
        ['作品上传', '能否成功上传并存证', '2026-03-27', '2026-03-27', '李文博'],
        ['管理员审核', '能否通过/拒绝', '2026-03-28', '2026-03-28', '李文博'],
        ['投票上链', '投票后是否写入区块', '2026-03-29', '2026-03-29', '李文博'],
        ['区块链浏览器', '区块/交易是否可查', '2026-03-30', '2026-03-30', '李文博'],
        ['哈希校验', '第三方能否验证', '2026-03-31', '2026-03-31', '李文博']
    ]
)

add_heading('3.5.2 集成测试计划', level=3)
add_table(
    ['项目', '内容'],
    [
        ['集成测试范围', '各功能模块 + 全链路'],
        ['集成测试方法', 'Playwright + 人工验证'],
        ['集成测试环境', 'Go 1.21.5, Node.js v24, MySQL 8.0.45, Chrome, Windows 11'],
        ['将产生的文档', '集成测试用例，测试报告'],
        ['集成测试任务/优先级', '见下表']
    ]
)

add_table(
    ['要测试的功能', '预期结果', '进度日期', '实际日期', '负责人'],
    [
        ['用户管理模块', '注册→登录→投票全流程', '2026-04-01', '2026-04-01', '李文博'],
        ['作品管理模块', '上传→审核→投票→排行', '2026-04-02', '2026-04-02', '李文博'],
        ['区块链模块', '作品存证+投票存证+链浏览器', '2026-04-03', '2026-04-03', '李文博'],
        ['第三方验证', '哈希校验页可用', '2026-04-04', '2026-04-04', '李文博']
    ]
)

add_heading('3.5.3 单元测试报告', level=3)
add_para('1. 用户管理单元测试')
add_table(
    ['模块名/函数名', '输入数据', '期望结果', '实际结果', '备注'],
    [
        ['register',
         'username: tester, password: test123, role: 0',
         '提示注册成功',
         '提示注册成功',
         '正常路径'],
        ['register',
         'username: admin(已存在), password: test123',
         '提示用户名已被占用',
         '提示用户名已被占用',
         '异常路径'],
        ['login',
         'username: admin, password: admin123',
         '登录成功，返回 token',
         '登录成功，返回 token',
         '正常路径'],
        ['login',
         'username: admin, password: WRONG',
         '提示用户名或密码错误',
         '提示用户名或密码错误',
         '异常路径'],
        ['HashPassword',
         'salt=abc, password=123456',
         '返回 64 字符十六进制',
         '返回 64 字符十六进制',
         '正常路径']
    ]
)

add_para('2. 区块链模块单元测试')
add_table(
    ['模块名/函数名', '输入数据', '期望结果', '实际结果', '备注'],
    [
        ['MerkleRoot',
         '[]string{} (空)',
         '返回 EMPTY_MERKLE_ROOT 哈希',
         '返回 EMPTY_MERKLE_ROOT 哈希',
         '边界'],
        ['MerkleRoot',
         '["abc"] (单元素)',
         '返回 abc 的 SHA-256',
         '返回一致',
         '边界'],
        ['MerkleRoot',
         '["a","b","c","d"]',
         '正确计算二叉树根',
         '与手算一致',
         '正常'],
        ['PackBlock',
         'index=1, prevHash=0000…0, 1 笔交易, difficulty=2',
         '返回哈希前缀为 "00" 的区块',
         '返回哈希前缀为 "00" 的区块',
         'PoW 验证'],
        ['PackBlock',
         'index=1, prevHash=0000…0, 0 笔交易',
         '返回创世区块，merkleRoot = EMPTY',
         '返回创世区块',
         '创世']
    ]
)

add_para('3. 作品管理单元测试')
add_table(
    ['模块名/函数名', '输入数据', '期望结果', '实际结果', '备注'],
    [
        ['SubmitPhoto',
         'title: 测试作品, imageHash: 64位十六进制, photographer_id: 100',
         '插入 photo + tx + 打包区块',
         '插入 photo + tx + 打包区块',
         '正常'],
        ['SubmitPhoto',
         'imageHash 已存在',
         '提示哈希冲突',
         '提示该图片已存在',
         '异常'],
        ['CastVote',
         'user: 101, photo: 1 (已审核, 非本人)',
         '插入 vote + tx + 打包区块',
         '插入成功 + 链上 +1',
         '正常'],
        ['CastVote',
         '同一用户对同一作品第二次投票',
         '提示已投过票',
         '提示已投过票',
         '幂等保护']
    ]
)

add_heading('3.5.4 集成测试报告', level=3)
add_para('集成测试覆盖 4 大模块，所有用例全部通过：')
add_table(
    ['功能点', '用例编号', '输入', '预期结果', '测试结果', '失败原因'],
    [
        ['用户管理', '002', 'username:tester password:test123',
         '注册账号成功', '符合预期', '无'],
        ['用户管理', '001', 'username:tester password:test123',
         '登录成功获取 token', '符合预期', '无'],
        ['作品管理', '003', '上传一张 800x600 PNG',
         '作品入库 + 上链', '符合预期', '无'],
        ['审核管理', '016', '对 photo_id=1 审核通过',
         '状态变为 1', '符合预期', '无'],
        ['投票管理', '004', '对 photo_id=1 投票',
         'vote_count +1 + 链上 +1', '符合预期', '无'],
        ['区块链', '017', '访问 /api/chain/state',
         '返回 latestIndex/totalTxs', '符合预期', '无'],
        ['区块链', '017', '访问 /api/chain/block/2',
         '返回区块详情 + 交易列表', '符合预期', '无'],
        ['存证校验', '005', '提交 photo.imageHash',
         '返回 photo + chainTx', '符合预期', '无']
    ]
)

add_pagebreak()

# ============== 4 设计总结 ==============
add_title('4  设计总结', level=1)

add_heading('4.1 部署手册', level=2)
add_para('本系统运行需要安装 MySQL 8.0、Node.js v18+、Go 1.21+。'
         '数据库与 Go 工具安装完成后，'
         '按照如下步骤即可完成安装部署工作：')

add_heading('4.1.1 区块链平台搭建', level=3)
add_para('本项目使用 Go 自研 PoW 区块链引擎，启动时自动创建创世区块：')
add_code('cd backend\ngo build -o wwmm-server.exe .\n./wwmm-server.exe\n# 服务启动后自动连接 MySQL 并创建创世区块 #1')

add_heading('4.1.2 数据库初始化', level=3)
add_para('（1）获取 sql/init.sql 文件。')
add_para('（2）打开 Navicat 或 MySQL CLI，连接数据库后执行 init.sql：')
add_code('mysql -h localhost -P 3306 -u root -p123456 < sql/init.sql')

add_para('（3）脚本会自动创建 wwmm_db 数据库、8 张数据表、2 个视图，'
         '并在 chain_state 中插入 1 条初始状态。')

add_heading('4.1.3 启动服务端', level=3)
add_para('（1）进入 backend 目录，编译：go build -o wwmm-server.exe .')
add_para('（2）直接运行：./wwmm-server.exe')
add_para('（3）看到 "[OK] 数据库连接成功" 和 "[OK] WWMM 后端启动于 :8080" 即为成功')

add_heading('4.1.4 启动客户端', level=3)
add_para('（1）进入 frontend 目录，安装依赖：npm install')
add_para('（2）启动开发服务器：npm run dev')
add_para('（3）浏览器访问 http://localhost:5173 即可使用平台')

add_heading('4.1.5 访问程序', level=3)
add_para('（1）在浏览器中输入 http://localhost:5173')
add_para('（2）页面将跳转至首页，展示热门作品和区块链状态')
add_para('（3）点击"登录"使用 demo 账号登录：')
add_bullets([
    '管理员：admin / admin123',
    '摄影师：photographer / photo123',
    '投票用户：voter / vote123'
])

add_heading('4.2 用户操作手册', level=2)

add_heading('4.2.1 用户登录', level=3)
add_para('输入网址进入平台登录页面，'
         '在账号输入框中输入自己的账号，'
         '在密码输入框中输入密码。')
add_para('如果不输入账号或密码，会提示错误。')
add_para('如果输入错误的账号或密码，会提示"用户名或密码错误"。')
add_para('如果输入正确的账号和密码，则显示登录成功，'
         '并跳转到首页。')

add_heading('4.2.2 用户注册', level=3)
add_para('点击右上角"免费注册"可跳转到注册界面。')
add_para('填写账号（3-16位字母数字下划线）、密码（至少6位）、'
         '真实姓名、手机号、邮箱、性别。')
add_para('选择注册为"普通用户"或"摄影师"。')
add_para('点击"注册并登录"后，提示注册成功并自动登录。')

add_heading('4.2.3 上传与投票功能', level=3)
add_para('摄影师：在"上传作品"页填写作品信息，拖拽或选择图片，点击"提交存证"。'
         '后端立即计算 SHA-256 哈希并打包上链。')
add_para('普通用户：在作品详情页点击"为它投一票"，'
         '后端写入 vote 表并打包上链，前端显示链上交易哈希。')
add_para('任意用户：在区块链浏览器中查看所有区块、交易，'
         '或通过哈希校验页输入 64 字符 SHA-256 验证作品归属。')

add_heading('4.3 总结感悟', level=2)
add_para('本次毕业设计《基于区块链的摄影作品投票存证系统》'
         '从选题到实现历时近 4 周，'
         '从最初的区块链入门，到自研 PoW 引擎的攻坚，'
         '再到 Solidity 合约的并行设计，'
         '整个过程让我对区块链技术的理解从"知道"变成了"会用"。')

add_para('技术上，我最大的收获是：区块链并不神秘，'
         '它的核心就是"哈希链 + 工作量证明 + 共识"。'
         '当我自己用 Go 写完一个 200 行的 PoW 引擎后，'
         '比特币白皮书里那些看似高深的设计，'
         '突然变得具体、生动、可验证。'
         '我深刻理解了为什么"不可篡改"是区块链的根本属性——'
         '因为要篡改一个区块，'
         '就要重做这个区块以及之后所有区块的工作量证明，'
         '成本高到不可接受。')

add_para('工程上，我学到了分层架构的重要性。'
         '本项目后端按 Controller / Service / DAO 分层，'
         '前端按 views / components / api 分层，'
         '配合 MVC 思路，让代码的可维护性和可测试性大幅提升。'
         '尤其在调试"投票是否真的上了链"这种端到端问题时，'
         '清晰的分层让定位问题的速度提高了数倍。')

add_para('业务上，我认识到区块链不是万能的，'
         '但它确实是解决"信任问题"的一把利器。'
         '当用户上传一张照片，系统立刻返回一段不可篡改的哈希上链，'
         '当用户投下一票，链上多了一笔不可伪造的交易，'
         '这些以前需要第三方公证才能实现的事情，'
         '现在通过一段代码就能完成。'
         '这是技术普惠的最好体现。')

add_para('最后，我要感谢汪铭杰老师和罗斌老师的悉心指导，'
         '感谢同学们在调试过程中提出的宝贵建议，'
         '感谢家人对我的支持。'
         '本项目虽然已经结题，但我对区块链技术应用的探索'
         '才刚刚开始。')

add_pagebreak()

# ============== 参考文献 ==============
add_title('参考文献', level=1)
add_para('[1] 钱可强. 机械制图[M]. 北京: 高等教育出版社, 2022.')
add_para('[2] 刘永刚, 袁建国, 刘思波. 深度剖析——硬盘固件级数据恢复[M]. 北京: 电子工业出版社, 2023.')
add_para('[3] 高曙明. 自动特征识别技术综述[J]. 计算机学报, 2019(3): 281-288.')
add_para('[4] Nakamoto S. Bitcoin: A Peer-to-Peer Electronic Cash System[R]. 2008.')
add_para('[5] Buterin V. A Next-Generation Smart Contract and Decentralized Application Platform[R]. Ethereum White Paper, 2014.')
add_para('[6] FISCO BCOS开源工作组. FISCO BCOS 技术白皮书 v2.0[R]. 2020.')
add_para('[7] 王元龙, 张亮, 陈尧. 区块链技术与应用[M]. 北京: 机械工业出版社, 2022.')
add_para('[8] 邵奇峰, 金澈清, 张召, 等. 区块链技术: 架构及进展[J]. 计算机学报, 2018, 41(5): 969-988.')
add_para('[9] Antonopoulos A M. Mastering Bitcoin[M]. 2nd Edition. O\'Reilly Media, 2017.')
add_para('[10] Solidity Documentation. https://docs.soliditylang.org/, 2024.')
add_para('[11] 邹均, 张海宁, 唐屹, 等. 区块链技术指南[M]. 北京: 机械工业出版社, 2016.')
add_para('[12] 袁勇, 王飞跃. 区块链技术发展现状与展望[J]. 自动化学报, 2016, 42(4): 481-494.')
add_para('[13] 华为区块链技术开发团队. 区块链技术及应用[M]. 北京: 清华大学出版社, 2019.')
add_para('[14] 李芳, 李赫. 基于区块链的数字内容版权保护方案研究[J]. 信息通信技术, 2019(3): 45-52.')
add_para('[15] 蔡亮, 李启雷, 梁秀波. 区块链技术进阶与实战[M]. 北京: 人民邮电出版社, 2020.')
add_para('[16] 阿里巴巴集团. 阿里云区块链服务开发指南[R]. 2021.')
add_para('[17] Hyperledger Foundation. Hyperledger Fabric Documentation v2.5[R]. 2023.')
add_para('[18] Wood G. Ethereum: A Secure Decentralised Generalised Transaction Ledger[R]. Ethereum Yellow Paper, 2014.')
add_para('[19] 张增骏, 董宁, 朱轩彤, 等. 深度探索区块链: Hyperledger技术与应用[M]. 北京: 机械工业出版社, 2018.')
add_para('[20] 余维, 刘琦, 杨晓宇, 等. 一种基于工作量证明的区块链共识机制优化算法[J]. 计算机应用研究, 2019, 36(10): 3070-3073.')

doc.save(OUT)
print(f'[OK] Saved to {OUT}')
print(f'[Size] {os.path.getsize(OUT)} bytes')
